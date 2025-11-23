"""
app/services/report_generation/docx_generator.py
Generación de documentos Word personalizados
"""

import os
from typing import Dict, List
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.models.student_profile import StudentProfile
from app.services.report_generation.data_visualizer import data_visualizer
from app.services.ai.gemini_service import gemini_service
from app.utils.logger import logger


class DOCXGenerator:
    """
    Generador de documentos Word personalizados
    Crea reportes detallados del perfil académico
    """
    
    def __init__(self):
        self.logger = logger
        # Directorio de salida
        self.output_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))),
            'generated', 'docx'
        )
        os.makedirs(self.output_dir, exist_ok=True)
    
    # =====================================================
    # MÉTODO PRINCIPAL
    # =====================================================
    
    def generate_detailed_report(
        self,
        profile: StudentProfile,
        user_id: int,
        report_type: str = 'completo'
    ) -> Dict:
        """
        Genera reporte detallado en Word
        
        Args:
            profile: Perfil del estudiante
            user_id: ID del usuario
            report_type: 'completo', 'ejecutivo', 'academico'
            
        Returns:
            Dict con ruta del archivo y metadata
        """
        try:
            self.logger.info(f"Generando DOCX para user_id={user_id}, tipo={report_type}")
            
            # Crear documento
            doc = Document()
            
            # Configurar estilos
            self._setup_styles(doc)
            
            # Construir documento según tipo
            if report_type == 'completo':
                self._build_complete_report(doc, profile, user_id)
            elif report_type == 'ejecutivo':
                self._build_executive_report(doc, profile, user_id)
            elif report_type == 'academico':
                self._build_academic_report(doc, profile, user_id)
            else:
                self._build_complete_report(doc, profile, user_id)
            
            # Guardar archivo
            filename = f"reporte_{report_type}_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)
            
            self.logger.info(f"DOCX generado: {filepath}")
            
            return {
                'success': True,
                'filepath': filepath,
                'filename': filename,
                'report_type': report_type,
                'file_size': os.path.getsize(filepath)
            }
            
        except Exception as e:
            self.logger.error(f"Error generando DOCX: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    # =====================================================
    # CONFIGURACIÓN DE ESTILOS
    # =====================================================
    
    def _setup_styles(self, doc: Document):
        """Configura estilos personalizados del documento"""
        try:
            # Estilo para título principal
            styles = doc.styles
            
            # Título 1 personalizado
            if 'Heading 1' in styles:
                h1 = styles['Heading 1']
                h1.font.size = Pt(24)
                h1.font.bold = True
                h1.font.color.rgb = RGBColor(41, 128, 185)
            
            # Título 2 personalizado
            if 'Heading 2' in styles:
                h2 = styles['Heading 2']
                h2.font.size = Pt(18)
                h2.font.bold = True
                h2.font.color.rgb = RGBColor(52, 152, 219)
            
            # Título 3 personalizado
            if 'Heading 3' in styles:
                h3 = styles['Heading 3']
                h3.font.size = Pt(14)
                h3.font.bold = True
                h3.font.color.rgb = RGBColor(100, 100, 100)
                
        except Exception as e:
            self.logger.warning(f"Error configurando estilos: {str(e)}")
    
    # =====================================================
    # TIPOS DE REPORTES
    # =====================================================
    
    def _build_complete_report(self, doc: Document, profile: StudentProfile, user_id: int):
        """Reporte completo con todas las secciones"""
        # Portada
        self._add_cover_page(doc, profile)
        doc.add_page_break()
        
        # Índice (placeholder)
        self._add_table_of_contents(doc)
        doc.add_page_break()
        
        # 1. Resumen Ejecutivo
        self._add_executive_summary_section(doc, profile)
        doc.add_page_break()
        
        # 2. Perfil del Estudiante
        self._add_student_profile_section(doc, profile)
        doc.add_page_break()
        
        # 3. Análisis Académico
        self._add_academic_analysis_section(doc, profile)
        doc.add_page_break()
        
        # 4. Preparación para Tesis
        self._add_thesis_readiness_section(doc, profile)
        doc.add_page_break()
        
        # 5. Fortalezas y Debilidades
        self._add_strengths_weaknesses_section(doc, profile)
        doc.add_page_break()
        
        # 6. Recomendaciones
        self._add_recommendations_section(doc, profile)
        doc.add_page_break()
        
        # 7. Plan de Acción
        self._add_action_plan_section(doc, profile, user_id)
    
    def _build_executive_report(self, doc: Document, profile: StudentProfile, user_id: int):
        """Reporte ejecutivo completo y seguro"""
        try:
            # ==========================================
            # 1. PORTADA
            # ==========================================
            from app.models.user import User
            user = User.query.get(user_id)
        
            title = doc.add_heading('Reporte de Rendimiento Académico', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
            doc.add_paragraph()
            doc.add_paragraph()
        
            # Nombre del estudiante
            if user:
                name = f"{user.first_name} {user.last_name}"
                name_para = doc.add_paragraph(name)
                name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                name_para.runs[0].font.size = Pt(20)
                name_para.runs[0].font.bold = True

                # Carrera si existe
                if hasattr(user, 'career') and user.career:
                    career_para = doc.add_paragraph(user.career)
                    career_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    career_para.runs[0].font.size = Pt(14)
        
            doc.add_paragraph()
            doc.add_paragraph()
        
            # Fecha
            date_para = doc.add_paragraph(datetime.now().strftime("%d de %B de %Y"))
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.runs[0].font.size = Pt(12)
            date_para.runs[0].italic = True
        
            doc.add_page_break()
        
            # ==========================================
            # 2. RESUMEN EJECUTIVO
            # ==========================================
            doc.add_heading('Resumen Ejecutivo', 1)
        
            # Usar resumen IA del perfil de forma segura
            summary = profile.ai_profile_summary
            if summary and summary.strip():
                # Dividir por saltos de línea y agregar párrafos
                lines = summary.split('\n')
                for line in lines:
                    text = line.strip()
                    if text:
                        para = doc.add_paragraph(text)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                doc.add_paragraph("Perfil académico en construcción. El estudiante está iniciando su proceso de análisis.")

            doc.add_page_break()
        
            # ==========================================
            # 3. MÉTRICAS CLAVE (TABLA)
            # ==========================================
            doc.add_heading('Métricas Principales', 1)
        
            stats = data_visualizer.get_summary_stats(profile, user_id)
        
            # Crear tabla
            metrics_table = doc.add_table(rows=7, cols=2)
            metrics_table.style = 'Light Shading Accent 1'
        
            # Datos de la tabla
            table_data = [
                ('Documentos Analizados', str(stats.get('total_documents', 0))),
                ('Sesiones Completadas', str(stats.get('total_sessions', 0))),
                ('Score Preparación Tesis', f"{stats.get('thesis_score', 0):.1f}/100"),
                ('Calidad de Escritura', f"{stats.get('writing_quality', 0):.1f}/100"),
                ('Vocabulario Técnico', f"{stats.get('vocabulary_score', 0):.1f}/100"),
                ('Span de Atención', f"{stats.get('attention_span', 0)} min"),
                ('Estilo de Aprendizaje', stats.get('learning_style', 'No determinado'))
            ]
        
            # Llenar tabla de forma segura
            for i, (label, value) in enumerate(table_data):
                row = metrics_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                # Bold en labels
                for para in row.cells[0].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
        
            doc.add_page_break()
        
            # ==========================================
            # 4. PREPARACIÓN PARA TESIS
            # ==========================================
            doc.add_heading('Preparación para Tesis', 1)
        
            thesis_score = float(profile.thesis_readiness_score or 0)
            level = profile.thesis_readiness_level or 'bajo'
        
            # Score principal
            score_para = doc.add_paragraph(f"Score: {thesis_score:.1f}/100")
            score_para.runs[0].font.size = Pt(18)
            score_para.runs[0].font.bold = True
        
            # Nivel
            level_text_map = {
                'bajo': 'Nivel Bajo - Requiere preparación significativa',
                'medio': 'Nivel Medio - Requiere fortalecimiento en áreas clave',
                'alto': 'Nivel Alto - Buena preparación con mejoras menores',
                'excelente': 'Nivel Excelente - Listo para comenzar'
            }
            doc.add_paragraph(level_text_map.get(level, 'No evaluado'))
        
            doc.add_paragraph()
        
            # Factores
            doc.add_heading('Factores Evaluados', 2)
            factors = [
                f"Documentos Analizados: {profile.total_documents_analyzed} (Óptimo: 10+)",
                f"Calidad de Escritura: {float(profile.avg_writing_quality or 0):.1f}/100 (Objetivo: 80+)",
                f"Vocabulario Técnico: {float(profile.avg_vocabulary_richness or 0):.1f}/100 (Objetivo: 75+)",
                f"Consistencia: {profile.writing_improvement_trend or 'Sin datos'}",
               f"Tiempo Estimado: {profile.estimated_preparation_months or 12} meses"
            ]
        
            for factor in factors:
                doc.add_paragraph(factor, style='List Bullet')
        
            doc.add_page_break()
        
            # ==========================================
            # 5. FORTALEZAS Y DEBILIDADES
            # ==========================================
            doc.add_heading('Fortalezas y Áreas de Oportunidad', 1)
        
            # Fortalezas
            doc.add_heading('Fortalezas Identificadas', 2)
        
            all_strengths = []
            if profile.academic_strengths:
                all_strengths.extend(profile.academic_strengths)
            if profile.writing_strengths:
                all_strengths.extend(profile.writing_strengths)
            if profile.technical_strengths:
                all_strengths.extend(profile.technical_strengths)
        
            if all_strengths:
                for strength in all_strengths[:5]:  # Máximo 5
                    doc.add_paragraph(strength, style='List Bullet')
            else:
                doc.add_paragraph("Aún construyendo perfil - fortalezas por identificar")
        
            doc.add_paragraph()
        
            # Debilidades
            doc.add_heading('Áreas de Oportunidad', 2)
        
            all_weaknesses = []
            if profile.academic_weaknesses:
                all_weaknesses.extend(profile.academic_weaknesses)
            if profile.writing_weaknesses:
                all_weaknesses.extend(profile.writing_weaknesses)
            if profile.areas_for_improvement:
                all_weaknesses.extend(profile.areas_for_improvement)
        
            if all_weaknesses:
                for weakness in all_weaknesses[:5]:  # Máximo 5
                    doc.add_paragraph(weakness, style='List Bullet')
            else:
                doc.add_paragraph("Sin áreas críticas identificadas")
        
            doc.add_page_break()
        
            # ==========================================
            # 6. RECOMENDACIONES
            # ==========================================
            doc.add_heading('Recomendaciones Personalizadas', 1)
        
            # Recomendaciones de estudio
            if profile.study_recommendations:
                doc.add_heading('Estrategias de Estudio', 2)
                for i, rec in enumerate(profile.study_recommendations[:5], 1):
                    doc.add_paragraph(f"{i}. {rec}")
        
            doc.add_paragraph()
        
            # Recursos recomendados
            if profile.resource_recommendations:
                doc.add_heading('Recursos Recomendados', 2)
                for resource in profile.resource_recommendations[:5]:
                    doc.add_paragraph(resource, style='List Bullet')
        
        except Exception as e:
            self.logger.error(f"Error en executive report: {str(e)}")
            import traceback
            traceback.print_exc()
            doc.add_paragraph("Error generando algunas secciones del reporte.")
    
    def _build_academic_report(self, doc: Document, profile: StudentProfile, user_id: int):
        """Reporte enfocado en aspectos académicos"""
        self._add_cover_page(doc, profile)
        doc.add_page_break()
        
        self._add_academic_analysis_section(doc, profile)
        doc.add_page_break()
        
        self._add_thesis_readiness_section(doc, profile)
        doc.add_page_break()
        
        self._add_strengths_weaknesses_section(doc, profile)
        doc.add_page_break()
        
        self._add_action_plan_section(doc, profile, user_id)
    
    # =====================================================
    # SECCIONES DEL DOCUMENTO
    # =====================================================
    
    def _add_cover_page(self, doc: Document, profile: StudentProfile):
        """Página de portada"""
        from app.models.user import User
        user = User.query.get(profile.user_id)
        
        # Título
        title = doc.add_heading('Reporte de Rendimiento Académico', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espacio
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Nombre del estudiante
        name = f"{user.first_name} {user.last_name}" if user else "Estudiante"
        name_para = doc.add_paragraph(name)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.runs[0].font.size = Pt(20)
        name_para.runs[0].font.bold = True
        
        # Carrera
        if user and hasattr(user, 'career'):
            career_para = doc.add_paragraph(user.career)
            career_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            career_para.runs[0].font.size = Pt(14)
        
        # Espacio
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Fecha
        date_para = doc.add_paragraph(datetime.now().strftime("%d de %B de %Y"))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(12)
        date_para.runs[0].italic = True
    
    def _add_table_of_contents(self, doc: Document):
        """Índice (placeholder)"""
        doc.add_heading('Índice', 1)
        
        toc_items = [
            '1. Resumen Ejecutivo',
            '2. Perfil del Estudiante',
            '3. Análisis Académico',
            '4. Preparación para Tesis',
            '5. Fortalezas y Debilidades',
            '6. Recomendaciones',
            '7. Plan de Acción'
        ]
        
        for item in toc_items:
            doc.add_paragraph(item, style='List Number')
    
    def _add_executive_summary_section(self, doc: Document, profile: StudentProfile):
        """Sección de resumen ejecutivo"""
        doc.add_heading('1. Resumen Ejecutivo', 1)
        
        # Usar resumen IA del perfil
        summary = profile.ai_profile_summary or "Perfil en construcción."
        
        paragraphs = summary.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_student_profile_section(self, doc: Document, profile: StudentProfile):
        """Sección de perfil del estudiante"""
        doc.add_heading('2. Perfil del Estudiante', 1)
        
        # Información básica
        doc.add_heading('2.1 Información General', 2)
        
        from app.models.user import User
        user = User.query.get(profile.user_id)
        
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        rows_data = [
            ('Nombre:', f"{user.first_name} {user.last_name}" if user else "N/A"),
            ('Carrera:', user.career if user and hasattr(user, 'career') else "N/A"),
            ('Ciclo Actual:', str(user.current_cycle) if user and hasattr(user, 'current_cycle') else "N/A"),
            ('Estilo de Aprendizaje:', profile.learning_style or "No determinado"),
            ('Última Actualización:', profile.last_updated.strftime('%d/%m/%Y') if profile.last_updated else "N/A")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    def _add_academic_analysis_section(self, doc: Document, profile: StudentProfile):
        """Sección de análisis académico"""
        doc.add_heading('3. Análisis Académico', 1)
        
        # Métricas principales
        doc.add_heading('3.1 Métricas Principales', 2)
        
        stats = data_visualizer.get_summary_stats(profile, profile.user_id)
        
        metrics_text = f"""
Documentos Analizados: {stats['total_documents']}
Sesiones Completadas: {stats['total_sessions']}
Calidad de Escritura: {stats['writing_quality']:.1f}/100
Riqueza de Vocabulario: {stats['vocabulary_score']:.1f}/100
Span de Atención: {stats['attention_span']} minutos
Tendencia de Mejora: {stats['improvement_trend']}
"""
        
        doc.add_paragraph(metrics_text.strip())
        
        # Interpretación
        doc.add_heading('3.2 Interpretación', 2)
        
        interpretation = self._generate_academic_interpretation(profile)
        doc.add_paragraph(interpretation)
    
    def _add_thesis_readiness_section(self, doc: Document, profile: StudentProfile):
        """Sección de preparación para tesis"""
        doc.add_heading('4. Preparación para Tesis', 1)
        
        # Score principal
        thesis_score = float(profile.thesis_readiness_score or 0)
        level = profile.thesis_readiness_level or 'bajo'
        
        doc.add_heading('4.1 Score de Preparación', 2)
        
        score_para = doc.add_paragraph(f"Score: {thesis_score:.1f}/100")
        score_para.runs[0].font.size = Pt(18)
        score_para.runs[0].font.bold = True
        
        level_text = {
            'bajo': 'Nivel Bajo - Requiere preparación significativa',
            'medio': 'Nivel Medio - Requiere fortalecimiento en áreas clave',
            'alto': 'Nivel Alto - Buena preparación con mejoras menores',
            'excelente': 'Nivel Excelente - Listo para comenzar'
        }
        
        doc.add_paragraph(level_text.get(level, 'No evaluado'))
        
        # Factores evaluados
        doc.add_heading('4.2 Factores Evaluados', 2)
        
        factors = [
            f"Documentos Analizados: {profile.total_documents_analyzed} (Óptimo: 10+)",
            f"Calidad de Escritura: {float(profile.avg_writing_quality or 0):.1f}/100 (Objetivo: 80+)",
            f"Vocabulario Técnico: {float(profile.avg_vocabulary_richness or 0):.1f}/100 (Objetivo: 75+)",
            f"Consistencia: {profile.writing_improvement_trend or 'Sin datos'}",
            f"Tiempo Estimado: {profile.estimated_preparation_months or 12} meses"
        ]
        
        for factor in factors:
            doc.add_paragraph(factor, style='List Bullet')
    
    def _add_strengths_weaknesses_section(self, doc: Document, profile: StudentProfile):
        """Sección de fortalezas y debilidades"""
        doc.add_heading('5. Fortalezas y Áreas de Oportunidad', 1)
        
        # Fortalezas
        doc.add_heading('5.1 Fortalezas Identificadas', 2)
        
        all_strengths = []
        if profile.academic_strengths:
            all_strengths.extend(profile.academic_strengths)
        if profile.writing_strengths:
            all_strengths.extend(profile.writing_strengths)
        if profile.technical_strengths:
            all_strengths.extend(profile.technical_strengths)
        
        if not all_strengths:
            doc.add_paragraph("Aún construyendo perfil - fortalezas por identificar")
        else:
            for strength in all_strengths:
                doc.add_paragraph(strength, style='List Bullet')
        
        # Debilidades
        doc.add_heading('5.2 Áreas de Oportunidad', 2)
        
        all_weaknesses = []
        if profile.academic_weaknesses:
            all_weaknesses.extend(profile.academic_weaknesses)
        if profile.writing_weaknesses:
            all_weaknesses.extend(profile.writing_weaknesses)
        if profile.areas_for_improvement:
            all_weaknesses.extend(profile.areas_for_improvement)
        
        if not all_weaknesses:
            doc.add_paragraph("Sin áreas críticas identificadas")
        else:
            for weakness in all_weaknesses:
                doc.add_paragraph(weakness, style='List Bullet')
    
    def _add_recommendations_section(self, doc: Document, profile: StudentProfile):
        """Sección de recomendaciones"""
        doc.add_heading('6. Recomendaciones Personalizadas', 1)
        
        # Recomendaciones de estudio
        if profile.study_recommendations:
            doc.add_heading('6.1 Estrategias de Estudio', 2)
            for rec in profile.study_recommendations:
                doc.add_paragraph(rec, style='List Number')
        
        # Recursos recomendados
        if profile.resource_recommendations:
            doc.add_heading('6.2 Recursos Recomendados', 2)
            for resource in profile.resource_recommendations:
                doc.add_paragraph(resource, style='List Bullet')
    
    def _add_action_plan_section(self, doc: Document, profile: StudentProfile, user_id: int):
        """Sección de plan de acción"""
        doc.add_heading('7. Plan de Acción', 1)
        
        doc.add_paragraph(
            "Basado en el análisis de tu perfil, te proponemos el siguiente plan de acción:"
        )
        
        # Generar plan con IA
        action_items = self._generate_action_plan_with_ai(profile, user_id)
        
        # Tabla de plan
        plan_table = doc.add_table(rows=len(action_items) + 1, cols=3)
        plan_table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = plan_table.rows[0].cells
        headers[0].text = 'Acción'
        headers[1].text = 'Prioridad'
        headers[2].text = 'Plazo'
        
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Acciones
        for i, action in enumerate(action_items, 1):
            cells = plan_table.rows[i].cells
            cells[0].text = action['action']
            cells[1].text = action['priority']
            cells[2].text = action['timeline']
    
    def _add_key_metrics_table(self, doc: Document, profile: StudentProfile):
        """Tabla de métricas clave"""
        doc.add_heading('Métricas Clave', 2)
        
        stats = data_visualizer.get_summary_stats(profile, profile.user_id)
        
        metrics_table = doc.add_table(rows=7, cols=2)
        metrics_table.style = 'Light Shading Accent 1'
        
        metrics_data = [
            ('Documentos Analizados', str(stats['total_documents'])),
            ('Sesiones Completadas', str(stats['total_sessions'])),
            ('Score Preparación Tesis', f"{stats['thesis_score']:.1f}/100"),
            ('Calidad de Escritura', f"{stats['writing_quality']:.1f}/100"),
            ('Vocabulario Técnico', f"{stats['vocabulary_score']:.1f}/100"),
            ('Span de Atención', f"{stats['attention_span']} min"),
            ('Estilo de Aprendizaje', stats['learning_style'])
        ]
        
        for i, (metric, value) in enumerate(metrics_data):
            metrics_table.rows[i].cells[0].text = metric
            metrics_table.rows[i].cells[1].text = value
            metrics_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    # =====================================================
    # HELPERS CON IA
    # =====================================================
    
    def _generate_academic_interpretation(self, profile: StudentProfile) -> str:
        """Genera interpretación académica con IA"""
        try:
            prompt = f"""Como asesor académico, genera una interpretación breve (150 palabras) del siguiente perfil:

Documentos: {profile.total_documents_analyzed}
Sesiones: {profile.total_sessions_completed}
Calidad escritura: {float(profile.avg_writing_quality or 0):.1f}/100
Vocabulario: {float(profile.avg_vocabulary_richness or 0):.1f}/100
Tendencia: {profile.writing_improvement_trend or 'sin datos'}

Escribe en párrafo, sin formato markdown."""
            
            response = gemini_service.generate_content(prompt, profile.user_id, "academic_interpretation")
            
            if response['success']:
                return response['content']
                
        except:
            pass
        
        return "El perfil muestra un estudiante en desarrollo académico. Se recomienda continuar con el análisis regular para obtener métricas más precisas."
    
    def _generate_action_plan_with_ai(self, profile: StudentProfile, user_id: int) -> List[Dict]:
        """Genera plan de acción con IA"""
        try:
            prompt = f"""Genera 5 acciones específicas para este perfil académico:

Thesis score: {float(profile.thesis_readiness_score or 0):.0f}/100
Nivel: {profile.thesis_readiness_level}
Documentos: {profile.total_documents_analyzed}

Para cada acción, especifica: acción|prioridad|plazo
Prioridades: Alta, Media, Baja
Plazos: Inmediato, 1 mes, 3 meses, 6 meses

Formato: accion|prioridad|plazo (una por línea)"""
            
            response = gemini_service.generate_content(prompt, user_id, "action_plan")
            
            if response['success']:
                lines = [l.strip() for l in response['content'].split('\n') if '|' in l]
                actions = []
                for line in lines[:5]:
                    parts = line.split('|')
                    if len(parts) == 3:
                        actions.append({
                            'action': parts[0].strip(),
                            'priority': parts[1].strip(),
                            'timeline': parts[2].strip()
                        })
                if actions:
                    return actions
                    
        except:
            pass
        
        # Fallback
        return [
            {'action': 'Incrementar documentos analizados', 'priority': 'Alta', 'timeline': '1 mes'},
            {'action': 'Completar 2 sesiones semanales', 'priority': 'Alta', 'timeline': 'Inmediato'},
            {'action': 'Mejorar calidad de escritura', 'priority': 'Media', 'timeline': '3 meses'},
            {'action': 'Ampliar vocabulario técnico', 'priority': 'Media', 'timeline': '3 meses'},
            {'action': 'Revisar progreso mensualmente', 'priority': 'Baja', 'timeline': '6 meses'}
        ]


# Instancia única
docx_generator = DOCXGenerator()